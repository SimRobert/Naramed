from pydoc import doc
from typing import Annotated, Any

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from tempfile import NamedTemporaryFile

import datetime
import os
import re
import shutil

import whisper
import language_tool_python

app = FastAPI()

SAVED_REPORTS_DIR = Path("saved_reports")
REPORT_TITLE = "Raport Ecocardiografic"
ECHO_DATA_TITLE = "Date ecografice:"
CONCLUSION_LABEL = "Concluzie"
CONCLUSION_TITLE = "Concluzie:"
THANK_YOU_TEXT = "Multumim ca ati apelat la serviciile noastre!"
LIST_BULLET_STYLE = "List Bullet"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
BRAND_BLUE = RGBColor(0x2E, 0x75, 0xB6)

ALLOWED_ORIGINS = ["http://localhost:5173"]

MEDICAL_UNITS = {
    "FE": "%",
    "TAP": "",
}

FormStr = Annotated[str, Form(...)]
UploadedAudio = Annotated[UploadFile, File(...)]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # sau ["*"] pentru teste
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Ruleaza whisper
model = whisper.load_model("medium") 

# Corectare text cu LanguageTool
tool = language_tool_python.LanguageTool('ro')

@app.delete(
    "/delete_report/",
    responses={404: {"description": "Fisierul nu exista"}}
)
def delete_report(patient: FormStr, filename: FormStr):
    file_path = SAVED_REPORTS_DIR / patient / filename

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Fisierul nu exista")

    file_path.unlink()
    return {"success": True, "message": "Fisier sters"}


@app.delete(
    "/delete_patient_folder/",
    responses={
        404: {"description": "Folderul nu exista"},
        500: {"description": "Eroare la stergere"}
    }
)
def delete_patient_folder(patient: FormStr):
    folder_path = SAVED_REPORTS_DIR / patient

    if not folder_path.exists():
        raise HTTPException(status_code=404, detail="Folderul nu exista")

    try:
        shutil.rmtree(folder_path)
        return {"success": True, "message": "Folder sters complet"}
    except OSError as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Eroare la stergere: {str(exc)}"
        ) from exc

@app.post("/update_report/")
async def update_report(
    patient_name: FormStr,
    report_name: FormStr,
    new_content: FormStr
):
    base_dir = "saved_reports"
    patient_dir = os.path.join(base_dir, patient_name)
    os.makedirs(patient_dir, exist_ok=True)

    # 1. Salveaza noul raport.txt
    txt_path = os.path.join(patient_dir, f"{report_name}.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(new_content)

    # 2. Parseaza informatia medicala
    data = parse_medical_data(new_content)

    # 3. Genereaza raport.docx
    doc = Document()

    # Titlu principal
    doc.add_heading(REPORT_TITLE, 0)

    # Subtitlu "Date ecografice"
    p1 = doc.add_paragraph()
    r1 = p1.add_run(ECHO_DATA_TITLE)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # Lista de date
    for label, entry in data.items():
        if label == "Concluzie":
            continue
        val = entry["value"]
        if isinstance(val, tuple):
            val = " x ".join(val)
        doc.add_paragraph(f"{label}: {val} mm", style="List Bullet")

    # Subtitlu si continut concluzie
    if "Concluzie" in data:
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        r2 = p2.add_run(CONCLUSION_TITLE)
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        doc.add_paragraph(data["Concluzie"]["value"])

    # Multumire
    doc.add_paragraph()
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.style = doc.styles["Normal"]
    r3 = p3.add_run("Multumim ca ati apelat la serviciile noastre!")
    r3.italic = True
    r3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # 4. Salvare DOCX
    docx_path = os.path.join(patient_dir, f"{report_name}.docx")
    doc.save(docx_path)

    return JSONResponse(content={"message": "Raport actualizat cu succes"})


@app.get("/list_all_reports")
def list_all_reports():
    base_path = Path("saved_reports")
    result = {}
    if not base_path.exists():
        return result
    for patient_folder in base_path.iterdir():
        if patient_folder.is_dir():
            file_list = []
            for f in patient_folder.glob("*"):
                if f.is_file():
                    timestamp = datetime.datetime.fromtimestamp(f.stat().st_mtime)
                    file_list.append({
                        "name": f.name,
                        "timestamp": timestamp.strftime("%d.%m.%Y %H:%M")
                    })
            result[patient_folder.name] = file_list
    return result 

@app.get("/download/{patient}/{filename}")
def download_file(patient: str, filename: str):
    file_path = Path(f"saved_reports/{patient}/{filename}")
    if file_path.exists():
        return FileResponse(file_path, filename=filename)
    return {"error": "File not found"}

@app.post("/transcribe/")
async def upload_audio(file: UploadedAudio):
    file_location = Path(f"temp_{file.filename}")

    with file_location.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        result = model.transcribe(str(file_location), language="ro", beam_size=5)
        matches = tool.check(result["text"])
        corrected_text = language_tool_python.utils.correct(result["text"], matches).strip()

        return {
            "original": result["text"],
            "corectat": corrected_text
        }
    finally:
        if file_location.exists():
            file_location.unlink()

@app.post("/generate_report/")
async def generate_report(
    transcript: FormStr,
    patient_name: FormStr
):
    # 1. Parseaza datele medicale
    data = parse_medical_data(transcript)

    # 2. Genereaza documentul .docx
    doc = Document()
    doc.add_heading("Raport Ecocardiografic", 0)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Date ecografice:")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # 3. Dictionar cu unitati
    units = {
        "FE": "%",
        "TAP": "",
    }

    # 4. Sorteaza campurile dupa ordinea in text (pozitie)
    sorted_items = sorted(
        [(k, v["value"], v["pos"]) for k, v in data.items() if k != "Concluzie"],
        key=lambda item: item[2]
    )

    for label, value, _ in sorted_items:
        if isinstance(value, tuple):
            value = " x ".join(value)
        unit = units.get(label, "mm")
        doc.add_paragraph(f"{label}: {val} mm", style=LIST_BULLET_STYLE)

    if "Concluzie" in data:
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Concluzie:")
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        doc.add_paragraph(data["Concluzie"]["value"])

    doc.add_paragraph()
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p3.add_run(THANK_YOU_TEXT)
    r3.italic = True
    r3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(
            tmp.name,
            media_type=DOCX_MEDIA_TYPE,
            filename=f"{patient_name}.docx"
        )



def parse_medical_data(text: str):
    data = {}

    # TAP complex tratat separat ca prioritar
    complex_tap_match = re.search(
        r"TAP\s+(\d+)[,\s]+(?:BAR|bar[ăa])[\s,]+(\d+)[,\s]+(?:BAR|bar[ăa])[\s,]+(\d+)",
        text,
        re.IGNORECASE
    )
    if complex_tap_match:
        data["TAP"] = {
            "value": "/".join(complex_tap_match.groups()),
            "pos": complex_tap_match.start()
        }

    patterns = {
        "Aorta la Inel": r"aorta\s+(?:la|în|in dreptul|în zona)\s+(?:inel(?:ul)?|segmentul)?\s*(\d+)",
        "Aorta la Sinusuri": r"aorta\s+(?:la|în)\s+sinusuri\s*(\d+)",
        "Aorta Ascendenta": r"aorta\s+ascendent[aă]\s*(\d+)",
        "AS": r"\bAS\s+(\d+)",
        "VD": r"\bVD\s+(\d+)",
        "SIV": r"\bSIV\s+(\d+)",
        "VS": r"\bVS\s+(\d+)x(\d+)",
        "Perete Posterior": r"perete\s+posterior\s*(\d+)",
        "FE": r"frac(?:tie|ție)(?: de ejec[tț]ie)?\s*(\d+)%?",
        "TAP": r"\bTAP\s+(\d+)",  
        "VMAX Aortica": r"valva\s+aortic[aă]\s+VMAX\s+(\d+[.,]\d+)",
        "VMAX Pulmonara": r"valva\s+pulmonar[aă]\s+VMAX\s+(\d+[.,]\d+)",
        "VMAX Tricuspidă": r"[Vv]alva\s+tr[iî]cuspid[aă]\s+VMAX\s+(\d+[.,]\d+)",
        "PSVD": r"\bPSVD\s+(\d+[.,]\d+)",
        "VMAX Arc Aortic": r"arc[aă]?\s+aortic\s+VMAX\s+(\d+[.,]\d+)",
        "Concluzie": r"[Cc]oncluzi[iî]?[ei][:,]?\s*(.+)",
    }

    for key, pattern in patterns.items():
        if key == "TAP" and "TAP" in data:
            continue
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.groups() if len(match.groups()) > 1 else match.group(1)
            data[key] = {"value": value, "pos": match.start()}

    return data

@app.post("/save_report/")
async def save_report(
    transcript: FormStr,
    patient_name: FormStr
):
    base_path = SAVED_REPORTS_DIR / patient_name

    if base_folder.exists():
        return JSONResponse(
            status_code=400,
            content={"error": "report_exists"}
        )

    base_folder.mkdir(parents=True, exist_ok=True)

    # 1. Salveaza .txt
    txt_path = base_folder / f"{patient_name}.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(transcript)

    # 2. Genereaza .docx
    data = parse_medical_data(transcript)
    doc = Document()

    doc.add_heading("Raport Ecocardiografic", 0)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Date ecografice:")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    for label, entry in data.items():
        if label == "Concluzie":
            continue
        val = entry["value"]
        if isinstance(val, tuple):
            val = " x ".join(val)
        doc.add_paragraph(f"{label}: {val} mm", style="List Bullet")


    if "Concluzie" in data:
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Concluzie:")
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        doc.add_paragraph(data["Concluzie"]["value"])


    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p3.add_run("Multumim ca ati apelat la serviciile noastre!")
    r3.italic = True
    r3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    docx_path = base_folder / f"{patient_name}.docx"
    doc.save(docx_path)

    return {"message": "Salvare cu succes", "folder": str(base_folder)}
